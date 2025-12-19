#!/usr/bin/env python3
"""
Jupiter CRM Audit V6-TEST
Professional CRM Data Quality Analysis Tool

Version: 6.0-TEST
Author: Stephanie Jupiter Jacca
Company: WBSE (We Bring Support & Expertise)
Contact: wbse.consult@gmail.com
Portfolio: https://stephaniejj.github.io

Features:
- DEMO mode (100 rows limit for free version)
- 6 advanced analysis functions (Cold Contacts, Churn Risk, Email Validity, etc.)
- 4 new business KPIs
- Chart legends for better UX
- Advanced Metrics section in PDF reports
- Power BI-style visualizations
- Bronze/Gold design theme
"""

# Jupiter CRM Audit V6-TEST

import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import io
import os
import base64
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==================== DEMO MODE CONFIGURATION ====================
DEMO_MODE = True  # Set to False for PRO version (unlimited data)
MAX_ROWS_DEMO = 100
SUPPORT_EMAIL = "wbse.consult@gmail.com"
PORTFOLIO_URL = "https://stephaniejj.github.io/#home"

def get_upgrade_message(total_rows, file_type):
    """Generate upgrade message for DEMO mode - NO PRICING"""
    return f"""
⚠️ **FREE VERSION - LIMITED TO {MAX_ROWS_DEMO} ROWS**

Your file contains **{total_rows:,} {file_type}**.  
Only the first **{MAX_ROWS_DEMO} will be analyzed** in the free version.

🚀 **Want to analyze all your data?**

Contact us for PRO access:
📧 {SUPPORT_EMAIL}  
🌐 {PORTFOLIO_URL}
"""

# ==================== PAGE CONFIGURATION ====================
st.set_page_config(
    page_title="Jupiter CRM Audit",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== CUSTOM CSS STYLING ====================
st.markdown("""
<style>
    /* Apple-inspired design */
    
    .stApp {
        background-color: #FFFFFF;
    }
    
    [data-testid="stSidebar"] {
        background-color: #86868B;
    }
    
    [data-testid="stSidebar"] * {
        color: #1D1D1F !important;
    }
    
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3 {
        color: #1D1D1F !important;
    }
    
    [data-testid="stFileUploader"] span,
    [data-testid="stFileUploader"] p,
    [data-testid="stFileUploader"] small {
        color: #1D1D1F !important;
    }
    
    :root {
        --primary-color: #CD7F32;
        --secondary-color: #B8860B;
        --accent-color: #DAA520;
        --background-color: #FFFFFF;
        --text-color: #1D1D1F;
        --secondary-text: #86868B;
    }

    .hero-section {
        text-align: center;
        padding: 3rem 0 2rem 0;
        background: #FFFFFF;
        margin-bottom: 2rem;
    }

    .hero-title {
        font-size: 3.5rem;
        font-weight: 600;
        color: #CD7F32;
        margin-bottom: 0.5rem;
        letter-spacing: -0.5px;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
    }

    .hero-subtitle {
        font-size: 1.3rem;
        color: #86868B;
        margin-bottom: 1rem;
        font-weight: 400;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Text", sans-serif;
    }

    .button-container {
        display: flex;
        justify-content: center;
        align-items: center;
        gap: 1.5rem;
        margin: 2.5rem 0;
        padding: 0;
    }

    .interactive-btn {
        display: inline-block;
        padding: 12px 32px;
        background-color: transparent;
        color: #CD7F32;
        border: 2px solid #E8E8ED;
        border-radius: 980px;
        text-decoration: none;
        font-size: 1rem;
        font-weight: 500;
        transition: all 0.3s ease;
        text-align: center;
        min-width: 120px;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Text", sans-serif;
        box-shadow: 0 0 10px rgba(205, 127, 50, 0.3);
    }

    .interactive-btn:hover {
        background-color: rgba(205, 127, 50, 0.1);
        color: #B8860B;
        border-color: #B8860B;
        transform: scale(1.02);
        box-shadow: 0 0 20px rgba(205, 127, 50, 0.5);
    }

    .footer {
        text-align: center;
        padding: 2rem 0;
        margin-top: 3rem;
        border-top: 1px solid #D2D2D7;
        color: #86868B;
        font-size: 0.9rem;
        font-weight: 400;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Text", sans-serif;
    }

    .stMetric {
        background-color: #F5F5F7;
        padding: 1.2rem;
        border-radius: 12px;
        border: 1px solid #E8E8ED;
    }
    
    .stMetric label {
        color: #1D1D1F !important;
        font-weight: 500 !important;
    }
    
    .stMetric [data-testid="stMetricValue"] {
        color: #1D1D1F !important;
        font-weight: 600 !important;
    }

    .stProgress > div > div > div > div {
        background-color: #CD7F32;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        gap: 2rem;
        background-color: transparent;
        border-bottom: 1px solid #D2D2D7;
    }
    
    .stTabs [data-baseweb="tab"] {
        color: #86868B;
        font-weight: 500;
    }
    
    .stTabs [aria-selected="true"] {
        color: #1D1D1F;
        border-bottom-color: #1D1D1F;
    }
    
    h1, h2, h3 {
        color: #1D1D1F !important;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif !important;
        font-weight: 600 !important;
    }
    
    p, span, div {
        color: #1D1D1F;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Text", sans-serif;
    }
    
    .streamlit-expanderHeader {
        background-color: #F5F5F7;
        color: #1D1D1F;
        border-radius: 8px;
        font-weight: 500;
    }

    /* Boutons Streamlit standards - Texte blanc */
    .stButton > button {
        color: #1D1D1F !important;
        background-color: #CD7F32 !important;
        border: 2px solid #E8E8ED !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Text", sans-serif !important;
        transition: all 0.3s ease !important;
    }
    
    .stButton > button:hover {
        background-color: #B8860B !important;
        border-color: #B8860B !important;
        box-shadow: 0 4px 12px rgba(205, 127, 50, 0.4) !important;
    }
    
    /* Boutons primary */
    .stButton > button[kind="primary"] {
        background-color: #CD7F32 !important;
        color: #1D1D1F !important;
    }


    /* Force white background for chart legends */
    div[style*="How to read this chart"] {
        background: #FFFFFF !important;
        color: #FFFFFF !important;
    }
    
    /* All divs inside markdown containing legends */
    .stMarkdown div {
        background-color: transparent !important;
    }
    
    /* Specific override for legend boxes */
    div[style*="padding: 20px"] {
        background: #FFFFFF !important;
    }


    /* File uploader text en blanc - REINFORCED */
    [data-testid="stFileUploader"] span,
    [data-testid="stFileUploader"] p,
    [data-testid="stFileUploader"] small,
    [data-testid="stFileUploader"] label,
    [data-testid="stFileUploader"] div {
        color: #FFFFFF !important;
    }
    
    /* Browse button text */
    [data-testid="stFileUploader"] button {
        color: #FFFFFF !important;
    }

</style>
""", unsafe_allow_html=True)

# ==================== SESSION STATE INITIALIZATION ====================
if 'contacts_df' not in st.session_state:
    st.session_state.contacts_df = None
if 'companies_df' not in st.session_state:
    st.session_state.companies_df = None
if 'tickets_df' not in st.session_state:
    st.session_state.tickets_df = None
if 'aggregated_df' not in st.session_state:
    st.session_state.aggregated_df = None
if 'pre_agg_scores' not in st.session_state:
    st.session_state.pre_agg_scores = None
if 'post_agg_score' not in st.session_state:
    st.session_state.post_agg_score = None
if 'audit_results' not in st.session_state:
    st.session_state.audit_results = None

# New V6 session state
if 'cold_analysis' not in st.session_state:
    st.session_state.cold_analysis = None
if 'email_analysis' not in st.session_state:
    st.session_state.email_analysis = None
if 'orphan_analysis' not in st.session_state:
    st.session_state.orphan_analysis = None
if 'ghost_companies' not in st.session_state:
    st.session_state.ghost_companies = None
if 'critical_tickets' not in st.session_state:
    st.session_state.critical_tickets = None
if 'churn_analysis' not in st.session_state:
    st.session_state.churn_analysis = None

if 'tickets_completeness' not in st.session_state:
    st.session_state.tickets_completeness = None
if 'companies_completeness' not in st.session_state:
    st.session_state.companies_completeness = None
if 'overall_quality' not in st.session_state:
    st.session_state.overall_quality = None
if 'quality_improvement' not in st.session_state:
    st.session_state.quality_improvement = None
if 'tickets_performance' not in st.session_state:
    st.session_state.tickets_performance = None
if 'top_industries' not in st.session_state:
    st.session_state.top_industries = None




# Jupiter CRM Audit V6-TEST

# ==================== UTILITY FUNCTIONS ====================

def load_data(file, file_type='data'):
    """Load data from uploaded CSV file with DEMO mode limit"""
    try:
        df = pd.read_csv(file)
        original_rows = len(df)
        
        # Apply DEMO mode limit
        if DEMO_MODE and original_rows > MAX_ROWS_DEMO:
            df = df.head(MAX_ROWS_DEMO)
            return df, original_rows, True  # is_limited=True
        
        return df, original_rows, False  # is_limited=False
        
    except Exception as e:
        st.error(f"❌ Error loading file: {str(e)}")
        return None, 0, False

def calculate_health_score(df, data_type='contacts'):
    """Calculate health score for a dataset"""
    if df is None or df.empty:
        return 0, []

    score = 100
    issues = []

    # Missing data penalty
    missing_pct = (df.isnull().sum().sum() / (len(df) * len(df.columns))) * 100
    if missing_pct > 0:
        penalty = min(missing_pct * 2, 30)
        score -= penalty
        issues.append(f"Missing data: {missing_pct:.1f}% (-{penalty:.1f} points)")

    # Duplicate penalty
    id_cols = [col for col in df.columns if 'id' in col.lower() or 'email' in col.lower()]
    if id_cols:
        dup_pct = (df.duplicated(subset=[id_cols[0]]).sum() / len(df)) * 100
        if dup_pct > 0:
            penalty = min(dup_pct * 3, 30)
            score -= penalty
            issues.append(f"Duplicates: {dup_pct:.1f}% (-{penalty:.1f} points)")

    # Empty fields penalty
    empty_fields = (df == '').sum().sum()
    if empty_fields > 0:
        empty_pct = (empty_fields / (len(df) * len(df.columns))) * 100
        penalty = min(empty_pct * 1.5, 20)
        score -= penalty
        issues.append(f"Empty fields: {empty_pct:.1f}% (-{penalty:.1f} points)")

    return max(score, 0), issues

def aggregate_data(contacts, companies, tickets):
    """Aggregate the three datasets"""
    if contacts is None or contacts.empty:
        return None

    result = contacts.copy()

    # Merge with companies
    if companies is not None and not companies.empty:
        company_id_col = next((col for col in contacts.columns if 'company' in col.lower() and 'id' in col.lower()), None)
        if company_id_col:
            company_main_id = next((col for col in companies.columns if 'id' in col.lower()), None)
            if company_main_id:
                # Convert both columns to string to avoid type mismatch
                result[company_id_col] = result[company_id_col].astype(str)
                companies = companies.copy()
                companies[company_main_id] = companies[company_main_id].astype(str)
                
                result = result.merge(
                    companies,
                    left_on=company_id_col,
                    right_on=company_main_id,
                    how='left',
                    suffixes=('', '_company')
                )

    # Add ticket statistics
    if tickets is not None and not tickets.empty:
        contact_id_col = next((col for col in contacts.columns if 'id' in col.lower()), None)
        ticket_contact_col = next((col for col in tickets.columns if 'contact' in col.lower() and 'id' in col.lower()), None)

        if contact_id_col and ticket_contact_col:
            # Convert both columns to string to avoid type mismatch
            result[contact_id_col] = result[contact_id_col].astype(str)
            tickets = tickets.copy()
            tickets[ticket_contact_col] = tickets[ticket_contact_col].astype(str)
            
            ticket_stats = tickets.groupby(ticket_contact_col).agg(
                ticket_count=('id', 'count') if 'id' in tickets.columns else (ticket_contact_col, 'count')
            ).reset_index()

            result = result.merge(
                ticket_stats,
                left_on=contact_id_col,
                right_on=ticket_contact_col,
                how='left'
            )
            result['ticket_count'] = result['ticket_count'].fillna(0)

    return result

def perform_audit(contacts, companies, tickets, aggregated):
    """Perform comprehensive audit analysis"""
    results = {
        'total_contacts': len(contacts) if contacts is not None else 0,
        'total_companies': len(companies) if companies is not None else 0,
        'total_tickets': len(tickets) if tickets is not None else 0,
        'duplicates': {},
        'missing_data': {},
        'data_quality': {},
        'recommendations': []
    }

    # Analyze duplicates
    if contacts is not None:
        email_col = next((col for col in contacts.columns if 'email' in col.lower()), None)
        if email_col:
            dup_count = contacts.duplicated(subset=[email_col]).sum()
            results['duplicates']['contacts'] = dup_count

    if companies is not None:
        name_col = next((col for col in companies.columns if 'name' in col.lower()), None)
        if name_col:
            dup_count = companies.duplicated(subset=[name_col]).sum()
            results['duplicates']['companies'] = dup_count

    # Analyze missing data
    if contacts is not None:
        results['missing_data']['contacts'] = contacts.isnull().sum().sum()
    if companies is not None:
        results['missing_data']['companies'] = companies.isnull().sum().sum()
    if tickets is not None:
        results['missing_data']['tickets'] = tickets.isnull().sum().sum()

    # Generate recommendations
    if results['duplicates'].get('contacts', 0) > 0:
        results['recommendations'].append({
            'priority': 'HIGH',
            'category': 'Data Cleaning',
            'issue': f"{results['duplicates']['contacts']} duplicate contacts found",
            'action': 'Implement automated deduplication process',
            'impact': 'Improve data accuracy and reduce confusion'
        })

    if results['missing_data'].get('contacts', 0) > len(contacts) * 0.1 if contacts is not None else False:
        results['recommendations'].append({
            'priority': 'MEDIUM',
            'category': 'Data Completeness',
            'issue': 'Significant missing data in contacts',
            'action': 'Implement data validation rules and mandatory fields',
            'impact': 'Enhance contact information quality'
        })

    return results

# ==================== V6 ANALYSIS FUNCTIONS ====================

def analyze_cold_contacts(df, days_threshold=90):
    """Analyse contacts froids (sans activité depuis X jours)"""
    if df is None or df.empty:
        return {'cold_count': 0, 'cold_pct': 0, 'total': 0}
    
    date_cols = [c for c in df.columns if 'last_activity' in c.lower() or 'last_contact' in c.lower()]
    
    if not date_cols:
        return {'cold_count': 0, 'cold_pct': 0, 'total': len(df), 'no_date_column': True}
    
    date_col = date_cols[0]
    
    try:
        df[date_col] = pd.to_datetime(df[date_col], errors='coerce')
        threshold_date = datetime.now() - timedelta(days=days_threshold)
        
        cold_mask = (df[date_col] < threshold_date) | (df[date_col].isna())
        cold_count = cold_mask.sum()
        
        return {
            'cold_count': int(cold_count),
            'cold_pct': round(cold_count / len(df) * 100, 1),
            'total': len(df),
            'threshold_days': days_threshold
        }
    except:
        return {'cold_count': 0, 'cold_pct': 0, 'total': len(df), 'error': True}


def analyze_email_validity(df):
    """Validation avancée des emails"""
    if df is None or df.empty:
        return {'valid': 0, 'invalid': 0, 'b2c': 0, 'total': 0}
    
    email_col = next((col for col in df.columns if 'email' in col.lower()), None)
    if not email_col:
        return {'valid': 0, 'invalid': 0, 'b2c': 0, 'total': 0}
    
    emails = df[email_col].dropna()
    
    email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    valid_syntax = emails.str.match(email_pattern, na=False)
    
    b2c_domains = ['gmail.com', 'yahoo.com', 'hotmail.com', 'outlook.com', 'live.com']
    b2c_mask = emails.str.lower().str.contains('|'.join(b2c_domains), na=False)
    
    return {
        'total': len(emails),
        'valid': int(valid_syntax.sum()),
        'invalid': int((~valid_syntax).sum()),
        'valid_pct': round(valid_syntax.sum() / len(emails) * 100, 1) if len(emails) > 0 else 0,
        'b2c_count': int(b2c_mask.sum()),
        'b2c_pct': round(b2c_mask.sum() / len(emails) * 100, 1) if len(emails) > 0 else 0
    }


def analyze_orphan_contacts(contacts_df):
    """Détecte contacts orphelins (sans company)"""
    if contacts_df is None or contacts_df.empty:
        return {'orphan_count': 0, 'orphan_pct': 0, 'total': 0}
    
    company_cols = [c for c in contacts_df.columns if 'company' in c.lower()]
    
    if not company_cols:
        return {'orphan_count': 0, 'orphan_pct': 0, 'total': len(contacts_df), 'no_company_column': True}
    
    company_col = company_cols[0]
    orphan_mask = contacts_df[company_col].isna() | (contacts_df[company_col] == '')
    orphan_count = orphan_mask.sum()
    
    return {
        'orphan_count': int(orphan_count),
        'orphan_pct': round(orphan_count / len(contacts_df) * 100, 1),
        'total': len(contacts_df)
    }


def analyze_companies_without_contacts(companies_df, contacts_df):
    """Détecte companies fantômes (sans contacts)"""
    if companies_df is None or companies_df.empty:
        return {'ghost_count': 0, 'ghost_pct': 0, 'total': 0}
    
    if contacts_df is None or contacts_df.empty:
        return {'ghost_count': len(companies_df), 'ghost_pct': 100, 'total': len(companies_df)}
    
    company_id_col = None
    for col in ['id', 'company_id', 'companyid']:
        if col in companies_df.columns:
            company_id_col = col
            break
    
    contact_company_col = None
    for col in ['company_id', 'companyid', 'company']:
        if col in contacts_df.columns:
            contact_company_col = col
            break
    
    if not company_id_col or not contact_company_col:
        return {'ghost_count': 0, 'ghost_pct': 0, 'total': len(companies_df), 'no_id_columns': True}
    
    companies_with_contacts = contacts_df[contact_company_col].dropna().unique()
    ghost_mask = ~companies_df[company_id_col].isin(companies_with_contacts)
    ghost_count = ghost_mask.sum()
    
    return {
        'ghost_count': int(ghost_count),
        'ghost_pct': round(ghost_count / len(companies_df) * 100, 1) if len(companies_df) > 0 else 0,
        'total': len(companies_df)
    }


def analyze_critical_tickets(tickets_df, hours_threshold=48):
    """Analyse tickets critiques ouverts >Xh"""
    if tickets_df is None or tickets_df.empty:
        return {'critical_count': 0, 'avg_resolution': 0, 'total': 0}
    
    date_col = None
    for col in ['created_date', 'createdate', 'created_at']:
        if col in tickets_df.columns:
            date_col = col
            break
    
    status_col = None
    for col in ['status', 'state', 'ticket_status']:
        if col in tickets_df.columns:
            status_col = col
            break
    
    if not date_col or not status_col:
        return {'critical_count': 0, 'avg_resolution': 0, 'total': len(tickets_df), 'no_required_columns': True}
    
    try:
        tickets_df[date_col] = pd.to_datetime(tickets_df[date_col], errors='coerce')
        
        open_statuses = ['open', 'new', 'pending', 'in progress', 'waiting']
        open_mask = tickets_df[status_col].str.lower().isin(open_statuses)
        
        threshold_date = datetime.now() - timedelta(hours=hours_threshold)
        critical_mask = open_mask & (tickets_df[date_col] < threshold_date)
        critical_count = critical_mask.sum()
        
        closed_col = None
        for col in ['closed_date', 'closedate', 'resolved_date']:
            if col in tickets_df.columns:
                closed_col = col
                break
        
        avg_resolution = 0
        if closed_col:
            tickets_df[closed_col] = pd.to_datetime(tickets_df[closed_col], errors='coerce')
            resolved = tickets_df[tickets_df[closed_col].notna()]
            if not resolved.empty:
                resolution_time = (resolved[closed_col] - resolved[date_col]).dt.total_seconds() / 3600
                avg_resolution = resolution_time.mean()
        
        return {
            'critical_count': int(critical_count),
            'total_open': int(open_mask.sum()),
            'total': len(tickets_df),
            'avg_resolution': round(avg_resolution, 1) if avg_resolution > 0 else 0,
            'threshold_hours': hours_threshold
        }
    except:
        return {'critical_count': 0, 'avg_resolution': 0, 'total': len(tickets_df), 'error': True}


def analyze_churn_risk(contacts_df, tickets_df=None):
    """Calcule score de risque churn par contact"""
    if contacts_df is None or contacts_df.empty:
        return {'at_risk_count': 0, 'at_risk_pct': 0, 'avg_score': 0, 'total': 0, 'arr_at_risk': 0}
    
    contacts_df = contacts_df.copy()
    contacts_df['churn_risk_score'] = 0
    
    # Signal 1: Inactivité
    date_cols = [c for c in contacts_df.columns if 'last_activity' in c.lower() or 'last_contact' in c.lower()]
    if date_cols:
        date_col = date_cols[0]
        try:
            contacts_df[date_col] = pd.to_datetime(contacts_df[date_col], errors='coerce')
            days_since = (datetime.now() - contacts_df[date_col]).dt.days
            contacts_df.loc[days_since > 90, 'churn_risk_score'] += 40
            contacts_df.loc[(days_since > 60) & (days_since <= 90), 'churn_risk_score'] += 20
            contacts_df.loc[(days_since > 30) & (days_since <= 60), 'churn_risk_score'] += 10
        except:
            pass
    
    # Signal 2: Email invalide
    email_col = next((col for col in contacts_df.columns if 'email' in col.lower()), None)
    if email_col:
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        invalid_email = ~contacts_df[email_col].str.match(email_pattern, na=False)
        contacts_df.loc[invalid_email, 'churn_risk_score'] += 15
    
    # Signal 3: Données incomplètes
    completeness = contacts_df.notna().sum(axis=1) / len(contacts_df.columns)
    contacts_df.loc[completeness < 0.5, 'churn_risk_score'] += 15
    
    at_risk_mask = contacts_df['churn_risk_score'] >= 70
    at_risk_count = at_risk_mask.sum()
    avg_score = contacts_df['churn_risk_score'].mean()
    
    arr_at_risk = 0
    arr_cols = [c for c in contacts_df.columns if c.lower() in ['arr', 'mrr', 'annual_revenue']]
    if arr_cols:
        arr_col = arr_cols[0]
        try:
            arr_at_risk = contacts_df.loc[at_risk_mask, arr_col].sum()
        except:
            pass
    
    return {
        'at_risk_count': int(at_risk_count),
        'at_risk_pct': round(at_risk_count / len(contacts_df) * 100, 1) if len(contacts_df) > 0 else 0,
        'avg_score': round(avg_score, 1),
        'total': len(contacts_df),
        'arr_at_risk': round(arr_at_risk, 0) if arr_at_risk > 0 else 0
    }



# Jupiter CRM Audit V6-TEST

# ==================== VISUALIZATION FUNCTIONS ====================

def add_chart_legend(legend_text):
    """Display chart legend using Streamlit native st.info()"""
    # Nettoyer le HTML du texte
    import re
    # Supprimer balises HTML
    clean_text = legend_text.replace('<br><br>', '\n\n')
    clean_text = clean_text.replace('<br>', '\n')
    clean_text = re.sub(r'<[^>]+>', '', clean_text)  # Enlever toutes les balises
    clean_text = clean_text.strip()
    
    # Ajouter titre
    full_text = "**ℹ️ How to read this chart**\n\n" + clean_text
    
    # Afficher avec st.info (natif Streamlit)
    st.info(full_text)
def analyze_tickets_completeness(tickets_df):
    """Calcule le taux de complétude des tickets"""
    if tickets_df is None or tickets_df.empty:
        return {'completeness_pct': 0, 'total_fields': 0, 'filled_fields': 0}
    
    total_cells = tickets_df.shape[0] * tickets_df.shape[1]
    filled_cells = total_cells - tickets_df.isnull().sum().sum()
    completeness_pct = (filled_cells / total_cells * 100) if total_cells > 0 else 0
    
    return {
        'completeness_pct': round(completeness_pct, 1),
        'total_fields': tickets_df.shape[1],
        'filled_fields': filled_cells,
        'total_cells': total_cells
    }


def analyze_companies_completeness(companies_df):
    """Calcule le taux de complétude des companies"""
    if companies_df is None or companies_df.empty:
        return {'completeness_pct': 0, 'total_fields': 0, 'filled_fields': 0}
    
    total_cells = companies_df.shape[0] * companies_df.shape[1]
    filled_cells = total_cells - companies_df.isnull().sum().sum()
    completeness_pct = (filled_cells / total_cells * 100) if total_cells > 0 else 0
    
    return {
        'completeness_pct': round(completeness_pct, 1),
        'total_fields': companies_df.shape[1],
        'filled_fields': filled_cells,
        'total_cells': total_cells
    }


def analyze_overall_quality(contacts_df, companies_df, tickets_df):
    """Calcule le score de qualité global"""
    if not any([contacts_df is not None, companies_df is not None, tickets_df is not None]):
        return {'overall_score': 0, 'breakdown': {}}
    
    scores = []
    breakdown = {}
    
    if contacts_df is not None and not contacts_df.empty:
        contact_score, _ = calculate_health_score(contacts_df, 'contacts')
        scores.append(contact_score)
        breakdown['contacts'] = round(contact_score, 1)
    
    if companies_df is not None and not companies_df.empty:
        company_score, _ = calculate_health_score(companies_df, 'companies')
        scores.append(company_score)
        breakdown['companies'] = round(company_score, 1)
    
    if tickets_df is not None and not tickets_df.empty:
        ticket_score, _ = calculate_health_score(tickets_df, 'tickets')
        scores.append(ticket_score)
        breakdown['tickets'] = round(ticket_score, 1)
    
    overall_score = sum(scores) / len(scores) if scores else 0
    
    return {
        'overall_score': round(overall_score, 1),
        'breakdown': breakdown
    }


def analyze_quality_improvement(pre_scores, post_score):
    """Calcule l'amélioration de qualité pré/post agrégation"""
    if not pre_scores or post_score is None:
        return {'improvement': 0, 'pre_avg': 0, 'post_score': 0}
    
    pre_values = [score for score, _ in pre_scores.values()]
    pre_avg = sum(pre_values) / len(pre_values) if pre_values else 0
    improvement = post_score - pre_avg
    
    return {
        'improvement': round(improvement, 1),
        'pre_avg': round(pre_avg, 1),
        'post_score': round(post_score, 1)
    }


def analyze_tickets_performance(tickets_df):
    """Analyse performance complète des tickets"""
    if tickets_df is None or tickets_df.empty:
        return {
            'open_count': 0,
            'closed_count': 0,
            'total_count': 0,
            'avg_resolution_hours': 0,
            'sla_compliance': None,
            'csat_score': None,
            'nps_score': None
        }
    
    status_col = None
    for col in ['status', 'state', 'ticket_status', 'hs_ticket_status']:
        if col in tickets_df.columns:
            status_col = col
            break
    
    open_count = 0
    closed_count = 0
    
    if status_col:
        open_statuses = ['open', 'new', 'pending', 'in progress', 'waiting']
        closed_statuses = ['closed', 'resolved', 'solved', 'completed']
        
        open_count = tickets_df[status_col].str.lower().isin(open_statuses).sum()
        closed_count = tickets_df[status_col].str.lower().isin(closed_statuses).sum()
    
    avg_resolution = 0
    created_col = None
    closed_col = None
    
    for col in ['created_date', 'createdate', 'created_at', 'hs_createdate']:
        if col in tickets_df.columns:
            created_col = col
            break
    
    for col in ['closed_date', 'closedate', 'resolved_date', 'hs_closed_date']:
        if col in tickets_df.columns:
            closed_col = col
            break
    
    if created_col and closed_col:
        try:
            tickets_df[created_col] = pd.to_datetime(tickets_df[created_col], errors='coerce')
            tickets_df[closed_col] = pd.to_datetime(tickets_df[closed_col], errors='coerce')
            
            resolved = tickets_df[tickets_df[closed_col].notna()].copy()
            if not resolved.empty:
                resolution_time = (resolved[closed_col] - resolved[created_col]).dt.total_seconds() / 3600
                avg_resolution = resolution_time.mean()
        except:
            pass
    
    sla_compliance = None
    sla_cols = ['sla_met', 'sla_status', 'within_sla', 'hs_sla_status']
    for col in sla_cols:
        if col in tickets_df.columns:
            try:
                met = tickets_df[col].notna().sum()
                total = len(tickets_df)
                sla_compliance = round(met / total * 100, 1) if total > 0 else 0
                break
            except:
                pass
    
    csat_score = None
    csat_cols = ['csat', 'customer_satisfaction', 'satisfaction_score', 'hs_csat']
    for col in csat_cols:
        if col in tickets_df.columns:
            try:
                csat_score = round(tickets_df[col].mean(), 1)
                break
            except:
                pass
    
    nps_score = None
    nps_cols = ['nps', 'net_promoter_score', 'nps_score', 'hs_nps']
    for col in nps_cols:
        if col in tickets_df.columns:
            try:
                nps_score = round(tickets_df[col].mean(), 1)
                break
            except:
                pass
    
    return {
        'open_count': int(open_count),
        'closed_count': int(closed_count),
        'total_count': len(tickets_df),
        'avg_resolution_hours': round(avg_resolution, 1) if avg_resolution > 0 else 0,
        'sla_compliance': sla_compliance,
        'csat_score': csat_score,
        'nps_score': nps_score
    }


def analyze_top_industries(companies_df, top_n=3):
    """Analyse les top industries"""
    if companies_df is None or companies_df.empty:
        return {'top_industries': [], 'total_companies': 0}
    
    industry_col = None
    for col in ['industry', 'sector', 'vertical', 'hs_industry']:
        if col in companies_df.columns:
            industry_col = col
            break
    
    if not industry_col:
        return {'top_industries': [], 'total_companies': len(companies_df), 'no_industry_column': True}
    
    industry_counts = companies_df[industry_col].value_counts().head(top_n)
    total = len(companies_df)
    
    top_industries = []
    for industry, count in industry_counts.items():
        if pd.notna(industry):
            percentage = round(count / total * 100, 1)
            top_industries.append({
                'name': str(industry),
                'count': int(count),
                'percentage': percentage
            })
    
    return {
        'top_industries': top_industries,
        'total_companies': total
    }


def create_powerbi_chart(fig, title):
    """Apply Apple-inspired styling with black text for readability"""
    fig.update_layout(
        title={
            'text': title,
            'font': {'size': 24, 'color': '#1D1D1F', 'family': 'SF Pro Display, sans-serif'},
            'x': 0.5,
            'xanchor': 'center'
        },
        plot_bgcolor='#FFFFFF',
        paper_bgcolor='#FFFFFF',
        font={'color': '#1D1D1F', 'family': 'SF Pro Text, sans-serif', 'size': 14},
        showlegend=True,
        legend={
            'bgcolor': '#FFFFFF',
            'bordercolor': '#D2D2D7',
            'borderwidth': 1,
            'font': {'color': '#1D1D1F', 'size': 13}
        },
        margin=dict(l=60, r=60, t=90, b=60),
        xaxis={'color': '#1D1D1F', 'gridcolor': '#E8E8ED'},
        yaxis={'color': '#1D1D1F', 'gridcolor': '#E8E8ED'}
    )
    
    # Update trace text colors to black
    fig.update_traces(textfont={'color': '#1D1D1F', 'size': 13})
    
    return fig


def plotly_fig_to_base64(fig, width=600, height=400):
    """Convert plotly figure to base64 for PDF embedding (requires kaleido)"""
    try:
        import plotly.io as pio
        img_bytes = fig.to_image(format="png", width=width, height=height)
        return base64.b64encode(img_bytes).decode()
    except ImportError:
        return None
    except Exception as e:
        return None


def generate_pdf_report(audit_results, pre_scores, post_score, 
                       cold_analysis=None, churn_analysis=None, 
                       critical_tickets=None, email_analysis=None,
                       orphan_analysis=None, ghost_companies=None):
    """Generate comprehensive PDF report with V6 Advanced Metrics"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#CD7F32'),
        spaceAfter=30,
        alignment=TA_CENTER
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#B8860B'),
        spaceAfter=12,
        spaceBefore=12
    )

    # Title
    story.append(Paragraph("Jupiter CRM Audit Report", title_style))
    story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    story.append(Spacer(1, 0.3*inch))

    # Executive Summary
    story.append(Paragraph("Executive Summary", heading_style))

    summary_data = [
        ['Metric', 'Value'],
        ['Total Contacts', f"{audit_results['total_contacts']:,}"],
        ['Total Companies', f"{audit_results['total_companies']:,}"],
        ['Total Tickets', f"{audit_results['total_tickets']:,}"],
        ['Post-Aggregation Score', f"{post_score:.1f}/100"]
    ]

    summary_table = Table(summary_data, colWidths=[3*inch, 3*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#CD7F32')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CD7F32'))
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.3*inch))

    # Advanced Metrics Analysis (V6)
    if any([cold_analysis, churn_analysis, critical_tickets, email_analysis, orphan_analysis, ghost_companies]):
        story.append(PageBreak())
        story.append(Paragraph("Advanced Business Metrics Analysis", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        advanced_data = [['Metric', 'Value', 'Status', 'Impact']]
        
        if cold_analysis and cold_analysis.get('cold_count', 0) > 0:
            status = '⚠️ Action Required' if cold_analysis['cold_pct'] > 30 else '✓ Acceptable'
            advanced_data.append([
                'Cold Contacts (>90d)',
                f"{cold_analysis['cold_pct']:.1f}% ({cold_analysis['cold_count']:,})",
                status,
                'Re-engagement needed'
            ])
        
        if email_analysis and email_analysis.get('total', 0) > 0:
            status = '✓ Good' if email_analysis['valid_pct'] > 90 else '⚠️ Needs Review'
            advanced_data.append([
                'Email Validity',
                f"{email_analysis['valid_pct']:.1f}% valid",
                status,
                f"{email_analysis['b2c_pct']:.1f}% B2C emails"
            ])
        
        if churn_analysis and churn_analysis.get('at_risk_count', 0) > 0:
            arr_text = f"${churn_analysis['arr_at_risk']:,.0f} ARR" if churn_analysis.get('arr_at_risk') else 'N/A'
            advanced_data.append([
                'Churn Risk',
                f"{churn_analysis['at_risk_count']} contacts",
                '🔴 Critical',
                arr_text
            ])
        
        if critical_tickets and critical_tickets.get('critical_count', 0) > 0:
            avg_text = f"{critical_tickets['avg_resolution']:.1f}h avg" if critical_tickets.get('avg_resolution') else 'N/A'
            advanced_data.append([
                'Critical Tickets (>48h)',
                f"{critical_tickets['critical_count']} open",
                '⚠️ Urgent',
                avg_text
            ])
        
        if orphan_analysis and orphan_analysis.get('orphan_count', 0) > 0:
            advanced_data.append([
                'Orphan Contacts',
                f"{orphan_analysis['orphan_pct']:.1f}% ({orphan_analysis['orphan_count']:,})",
                '⚠️ Data Quality',
                'No company association'
            ])
        
        if ghost_companies and ghost_companies.get('ghost_count', 0) > 0:
            advanced_data.append([
                'Ghost Companies',
                f"{ghost_companies['ghost_pct']:.1f}% ({ghost_companies['ghost_count']:,})",
                '⚠️ Cleanup Needed',
                'No contacts linked'
            ])
        
        advanced_table = Table(advanced_data, colWidths=[2*inch, 1.5*inch, 1.5*inch, 1.5*inch])
        advanced_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#CD7F32')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CD7F32')),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        story.append(advanced_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Immediate Actions Required
        story.append(Paragraph("Immediate Actions Required", heading_style))
        
        actions = []
        if cold_analysis and cold_analysis.get('cold_pct', 0) > 30:
            actions.append(f"Launch re-engagement campaign for {cold_analysis['cold_count']:,} cold contacts")
        
        if churn_analysis and churn_analysis.get('at_risk_count', 0) > 0:
            actions.append(f"Immediate outreach to {churn_analysis['at_risk_count']} high-risk contacts")
        
        if critical_tickets and critical_tickets.get('critical_count', 0) > 5:
            actions.append(f"Resolve {critical_tickets['critical_count']} critical tickets within 24h")
        
        if email_analysis and email_analysis.get('b2c_pct', 0) > 20:
            actions.append(f"Audit {email_analysis['b2c_pct']:.1f}% B2C emails (may need company emails)")
        
        if orphan_analysis and orphan_analysis.get('orphan_count', 0) > 0:
            actions.append(f"Link {orphan_analysis['orphan_count']:,} orphan contacts to companies")
        
        for action in actions:
            story.append(Paragraph(action, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

    # Recommendations
    story.append(PageBreak())
    story.append(Paragraph("Key Recommendations", heading_style))

    for rec in audit_results['recommendations']:
        story.append(Paragraph(f"<b>[{rec['priority']}] {rec['category']}</b>", styles['Normal']))
        story.append(Paragraph(f"Issue: {rec['issue']}", styles['Normal']))
        story.append(Paragraph(f"Action: {rec['action']}", styles['Normal']))
        story.append(Paragraph(f"Impact: {rec['impact']}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))

    # Footer
    story.append(Spacer(1, 0.5*inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#B8860B'),
        alignment=TA_CENTER
    )
    story.append(Paragraph("CRM Data Quality Analysis Report by Stephanie Jupiter Jacca from WBSE", footer_style))

    doc.build(story)
    buffer.seek(0)
    return buffer



# Jupiter CRM Audit V6-TEST

def generate_recommendations_document(audit_results, pre_scores, post_score,
                                     cold_analysis=None, churn_analysis=None,
                                     email_analysis=None, critical_tickets=None):
    """Generate detailed recommendations document (DOCX) saved to Desktop"""
    doc = Document()

    # Set document styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Jupiter CRM Audit - Strategic Recommendations & Consulting Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(205, 127, 50)

    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Prepared by: Stephanie Jupiter Jacca - WBSE")
    doc.add_paragraph("_" * 80)

    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    summary = doc.add_paragraph()
    summary.add_run(f"This comprehensive CRM audit analyzed {audit_results['total_contacts']:,} contacts, ")
    summary.add_run(f"{audit_results['total_companies']:,} companies, and {audit_results['total_tickets']:,} tickets. ")
    summary.add_run(f"The overall data quality score post-aggregation is {post_score:.1f}/100.\n\n")

    # V6 Advanced Metrics Section
    if any([cold_analysis, churn_analysis, email_analysis, critical_tickets]):
        doc.add_heading('Advanced Metrics Analysis', 1)
        
        if cold_analysis and cold_analysis.get('cold_count', 0) > 0:
            doc.add_heading('Inactive Contacts Alert', 2)
            doc.add_paragraph(
                f"{cold_analysis['cold_pct']:.1f}% of contacts ({cold_analysis['cold_count']:,}) "
                f"have not been active in the last {cold_analysis.get('threshold_days', 90)} days. "
                f"This represents a significant portion of your database that may require re-engagement campaigns."
            )
        
        if churn_analysis and churn_analysis.get('at_risk_count', 0) > 0:
            doc.add_heading('Churn Risk Assessment', 2)
            doc.add_paragraph(
                f"{churn_analysis['at_risk_count']} contacts identified as high churn risk (score ≥70/100). "
                f"Average churn risk score across all contacts: {churn_analysis.get('avg_score', 0):.1f}/100."
            )
            if churn_analysis.get('arr_at_risk', 0) > 0:
                doc.add_paragraph(
                    f"💰 Total ARR at risk: ${churn_analysis['arr_at_risk']:,.0f}",
                    style='List Bullet'
                )
        
        if email_analysis and email_analysis.get('total', 0) > 0:
            doc.add_heading('Email Validity Analysis', 2)
            doc.add_paragraph(
                f"{email_analysis['valid_pct']:.1f}% of emails are syntactically valid. "
                f"{email_analysis['b2c_pct']:.1f}% are personal B2C emails (Gmail, Yahoo, Hotmail, etc.), "
                f"which may indicate data quality issues in a B2B context."
            )
        
        if critical_tickets and critical_tickets.get('critical_count', 0) > 0:
            doc.add_heading('Support Performance Alert', 2)
            doc.add_paragraph(
                f"{critical_tickets['critical_count']} critical tickets have been open for more than "
                f"{critical_tickets.get('threshold_hours', 48)} hours. "
            )
            if critical_tickets.get('avg_resolution', 0) > 0:
                doc.add_paragraph(
                    f"Average resolution time: {critical_tickets['avg_resolution']:.1f} hours.",
                    style='List Bullet'
                )

    # Pre-Aggregation Scores
    doc.add_heading('Pre-Aggregation Health Scores', 2)
    if pre_scores:
        for obj_type, (score, issues) in pre_scores.items():
            p = doc.add_paragraph()
            p.add_run(f"{obj_type.capitalize()}: {score:.1f}/100\n").bold = True
            for issue in issues:
                doc.add_paragraph(f"  {issue}", style='List Bullet')

    # Key Findings
    doc.add_heading('Key Findings', 1)

    # Duplicates
    if audit_results['duplicates']:
        doc.add_heading('Duplicate Records', 2)
        for obj_type, count in audit_results['duplicates'].items():
            doc.add_paragraph(f"{obj_type.capitalize()}: {count:,} duplicates detected", style='List Bullet')

    # Missing Data
    if audit_results['missing_data']:
        doc.add_heading('Missing Data Analysis', 2)
        for obj_type, count in audit_results['missing_data'].items():
            doc.add_paragraph(f"{obj_type.capitalize()}: {count:,} missing values", style='List Bullet')

    # Strategic Recommendations
    doc.add_heading('Strategic Recommendations', 1)

    priority_order = {'HIGH': 1, 'MEDIUM': 2, 'LOW': 3}
    sorted_recs = sorted(audit_results['recommendations'], key=lambda x: priority_order.get(x['priority'], 4))

    for rec in sorted_recs:
        doc.add_heading(f"[{rec['priority']}] {rec['category']}", 2)
        doc.add_paragraph(f"Issue: {rec['issue']}")
        doc.add_paragraph(f"Recommended Action: {rec['action']}")
        doc.add_paragraph(f"Expected Impact: {rec['impact']}")
        doc.add_paragraph()

    # Footer
    doc.add_paragraph("_" * 80)
    footer = doc.add_paragraph()
    footer.add_run('\nCRM Data Quality Analysis Report by Stephanie Jupiter Jacca from WBSE\n').italic = True
    footer.add_run(f'Contact: wbse.consult@gmail.com\n').italic = True
    footer.add_run(f'Portfolio: https://stephaniejj.github.io/#home').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to Desktop
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
    filename = f'Jupiter_CRM_Recommendations_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    filepath = os.path.join(desktop_path, filename)

    try:
        doc.save(filepath)
        return filepath, True
    except Exception as e:
        return str(e), False



# Jupiter CRM Audit V6-TEST  

# ==================== SIDEBAR ====================
with st.sidebar:
    # Display logo
    try:
        st.image("LogoV6grey.png", use_column_width=True)
    except:
        st.warning("⚠️ Logo file 'LogoV6grey.png' not found in folder.")

    st.markdown("---")
    st.title("📂 Upload CRM Exports")

    st.markdown("### Step 1: Upload Required Files")

    contacts_file = st.file_uploader(
        "📧 Contacts CSV",
        type=['csv'],
        help="Upload your contacts export"
    )

    companies_file = st.file_uploader(
        "🏢 Companies CSV",
        type=['csv'],
        help="Upload your companies export"
    )

    tickets_file = st.file_uploader(
        "🎫 Tickets CSV",
        type=['csv'],
        help="Upload your tickets export"
    )

    # Load files into session state with DEMO mode handling
    if contacts_file:
        df, total_rows, is_limited = load_data(contacts_file, 'contacts')
        st.session_state.contacts_df = df
        
        if df is not None:
            if is_limited:
                st.warning(get_upgrade_message(total_rows, 'contacts'))
                st.info(f"📊 Analyzing first {MAX_ROWS_DEMO} rows (out of {total_rows:,})")
            st.success(f"✅ Contacts: {len(df):,} rows loaded")

    if companies_file:
        df, total_rows, is_limited = load_data(companies_file, 'companies')
        st.session_state.companies_df = df
        
        if df is not None:
            if is_limited:
                st.warning(get_upgrade_message(total_rows, 'companies'))
                st.info(f"📊 Analyzing first {MAX_ROWS_DEMO} rows (out of {total_rows:,})")
            st.success(f"✅ Companies: {len(df):,} rows loaded")

    if tickets_file:
        df, total_rows, is_limited = load_data(tickets_file, 'tickets')
        st.session_state.tickets_df = df
        
        if df is not None:
            if is_limited:
                st.warning(get_upgrade_message(total_rows, 'tickets'))
                st.info(f"📊 Analyzing first {MAX_ROWS_DEMO} rows (out of {total_rows:,})")
            st.success(f"✅ Tickets: {len(df):,} rows loaded")

# ==================== HERO SECTION ====================
st.markdown("""
<div class="hero-section">
    <div class="hero-title">Jupiter CRM Audit</div>
    <div class="hero-subtitle">Ai-driven Advanced Data Quality Analysis & Insights</div>
</div>
""", unsafe_allow_html=True)

# ==================== INTERACTIVE BUTTONS ====================
st.markdown("""
<div class="button-container">
    <a href="mailto:wbse.consult@gmail.com" class="interactive-btn" target="_blank">CONTACT</a>
    <a href="https://stephaniejj.github.io/#jupiter-apps" class="interactive-btn" target="_blank">AUDIT</a>
    <a href="https://stephaniejj.github.io/#home" class="interactive-btn" target="_blank">PORTFOLIO</a>
</div>
""", unsafe_allow_html=True)

# ==================== MAIN APPLICATION ====================

# Check if all three files are uploaded
all_files_uploaded = (
    st.session_state.contacts_df is not None and
    st.session_state.companies_df is not None and
    st.session_state.tickets_df is not None
)

if not all_files_uploaded:
    st.info("👆 Please upload all three required files (Contacts, Companies, Tickets) to begin the audit")
else:
    st.success("✅ All required files uploaded successfully!")

    # STEP 2: Pre-Aggregation Health Scores
    st.markdown("---")
    st.header("📊 Step 2: Pre-Aggregation Health Scores")

    if st.button("🔍 Calculate Pre-Aggregation Scores"):
        with st.spinner("Calculating health scores..."):
            progress_bar = st.progress(0)

            contacts_score, contacts_issues = calculate_health_score(st.session_state.contacts_df, 'contacts')
            progress_bar.progress(33)

            companies_score, companies_issues = calculate_health_score(st.session_state.companies_df, 'companies')
            progress_bar.progress(66)

            tickets_score, tickets_issues = calculate_health_score(st.session_state.tickets_df, 'tickets')
            progress_bar.progress(100)

            st.session_state.pre_agg_scores = {
                'contacts': (contacts_score, contacts_issues),
                'companies': (companies_score, companies_issues),
                'tickets': (tickets_score, tickets_issues)
            }

            st.success("✅ Pre-aggregation scores calculated!")

    if st.session_state.pre_agg_scores:
        col1, col2, col3 = st.columns(3)

        with col1:
            score, issues = st.session_state.pre_agg_scores['contacts']
            st.metric("📧 Contacts Health Score", f"{score:.1f}/100")
            with st.expander("View Issues"):
                for issue in issues:
                    st.write(f"{issue}")

        with col2:
            score, issues = st.session_state.pre_agg_scores['companies']
            st.metric("🏢 Companies Health Score", f"{score:.1f}/100")
            with st.expander("View Issues"):
                for issue in issues:
                    st.write(f"{issue}")

        with col3:
            score, issues = st.session_state.pre_agg_scores['tickets']
            st.metric("🎫 Tickets Health Score", f"{score:.1f}/100")
            with st.expander("View Issues"):
                for issue in issues:
                    st.write(f"{issue}")

    # STEP 3: Data Aggregation
    st.markdown("---")
    st.header("🔗 Step 3: Data Aggregation Tool")

    if st.button("🔄 Aggregate Data"):
        with st.spinner("Aggregating data..."):
            progress_bar = st.progress(0)

            st.session_state.aggregated_df = aggregate_data(
                st.session_state.contacts_df,
                st.session_state.companies_df,
                st.session_state.tickets_df
            )
            progress_bar.progress(100)

            if st.session_state.aggregated_df is not None:
                st.success(f"✅ Data aggregated successfully! Total records: {len(st.session_state.aggregated_df):,}")
            else:
                st.error("❌ Failed to aggregate data")

    if st.session_state.aggregated_df is not None:
        st.subheader("📋 Aggregated Data Preview")
        st.dataframe(st.session_state.aggregated_df.head(100), use_container_width=True)

    # STEP 4: Post-Aggregation Health Score
    st.markdown("---")
    st.header("📈 Step 4: Post-Aggregation Health Score")

    if st.session_state.aggregated_df is not None and st.button("📊 Calculate Post-Aggregation Score"):
        with st.spinner("Calculating post-aggregation score..."):
            score, issues = calculate_health_score(st.session_state.aggregated_df, 'aggregated')
            st.session_state.post_agg_score = (score, issues)
            st.success("✅ Post-aggregation score calculated!")

    if st.session_state.post_agg_score:
        score, issues = st.session_state.post_agg_score

        col1, col2 = st.columns([1, 2])
        with col1:
            # Gauge chart for score
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=score,
                domain={'x': [0, 1], 'y': [0, 1]},
                title={'text': "Overall Health Score", 'font': {'size': 22, 'color': '#1D1D1F'}},
                gauge={
                    'axis': {'range': [None, 100], 'tickwidth': 1, 'tickcolor': "#CD7F32"},
                    'bar': {'color': "#CD7F32"},
                    'bgcolor': "white",
                    'borderwidth': 2,
                    'bordercolor': "#B8860B",
                    'steps': [
                        {'range': [0, 50], 'color': '#FFE4B5'},
                        {'range': [50, 75], 'color': '#F4A460'},
                        {'range': [75, 100], 'color': '#DAA520'}
                    ],
                    'threshold': {
                        'line': {'color': "red", 'width': 4},
                        'thickness': 0.75,
                        'value': 90
                    }
                }
            ))
            fig.update_layout(
                paper_bgcolor='#FFFFFF',
                plot_bgcolor='#FFFFFF',
                font={'color': '#1D1D1F'},
                height=300
            )
            st.plotly_chart(fig, use_container_width=True)

        with col2:
            st.subheader("Identified Issues")
            for issue in issues:
                st.warning(f"⚠️ {issue}")

        # Comparison chart
        if st.session_state.pre_agg_scores:
            st.subheader("📊 Score Comparison: Pre vs Post Aggregation")

            pre_scores = [
                st.session_state.pre_agg_scores['contacts'][0],
                st.session_state.pre_agg_scores['companies'][0],
                st.session_state.pre_agg_scores['tickets'][0]
            ]
            avg_pre_score = sum(pre_scores) / len(pre_scores)

            comparison_data = pd.DataFrame({
                'Stage': ['Pre-Aggregation (Avg)', 'Post-Aggregation'],
                'Score': [avg_pre_score, score]
            })

            fig = px.bar(
                comparison_data,
                x='Stage',
                y='Score',
                title='Health Score Comparison',
                color='Score',
                color_continuous_scale=['#8B0000', '#CD7F32', '#DAA520'],
                text='Score'
            )
            fig.update_traces(texttemplate='%{text:.1f}', textposition='outside')
            fig = create_powerbi_chart(fig, 'Health Score Comparison')
            st.plotly_chart(fig, use_container_width=True)
            
            # Add legend
            add_chart_legend("""
            Compare your data quality BEFORE and AFTER merging.
            <br><br>
            <span style="color: white;">🔵 BEFORE MERGE:</span><br>
            Quality of each file separately
            <br><br>
            <span style="color: white;">🟢 AFTER MERGE:</span><br>
            Quality once all files combined
            <br><br>
            <span style="color: white;">🎯 TARGET:</span><br>
            Score above 80 = Good quality<br>
            Below 80 = Needs improvement
            """)

    # STEP 5: Launch Audit
    st.markdown("---")
    st.header("🚀 Step 5: Launch Complete Audit")

    if st.session_state.aggregated_df is not None and st.button("🔍 Launch Audit", type="primary"):
        with st.spinner("Performing comprehensive audit..."):
            progress_bar = st.progress(0)

            # Perform audit
            st.session_state.audit_results = perform_audit(
                st.session_state.contacts_df,
                st.session_state.companies_df,
                st.session_state.tickets_df,
                st.session_state.aggregated_df
            )
            progress_bar.progress(50)
            
            # V6 ANALYSES
            st.session_state.cold_analysis = analyze_cold_contacts(
                st.session_state.contacts_df,
                days_threshold=90
            )
            
            st.session_state.email_analysis = analyze_email_validity(
                st.session_state.contacts_df
            )
            
            st.session_state.orphan_analysis = analyze_orphan_contacts(
                st.session_state.contacts_df
            )
            
            st.session_state.ghost_companies = analyze_companies_without_contacts(
                st.session_state.companies_df,
                st.session_state.contacts_df
            )
            
            st.session_state.critical_tickets = analyze_critical_tickets(
                st.session_state.tickets_df,
                hours_threshold=48
            )
            
            st.session_state.churn_analysis = analyze_churn_risk(
                st.session_state.contacts_df,
                st.session_state.tickets_df
            )
            
            
            
            # PERFORMANCE METRICS ANALYSES
            st.session_state.tickets_completeness = analyze_tickets_completeness(
                st.session_state.tickets_df
            )
            
            st.session_state.companies_completeness = analyze_companies_completeness(
                st.session_state.companies_df
            )
            
            st.session_state.overall_quality = analyze_overall_quality(
                st.session_state.contacts_df,
                st.session_state.companies_df,
                st.session_state.tickets_df
            )
            
            st.session_state.quality_improvement = analyze_quality_improvement(
                st.session_state.pre_agg_scores,
                st.session_state.post_agg_score[0] if st.session_state.post_agg_score else None
            )
            
            st.session_state.tickets_performance = analyze_tickets_performance(
                st.session_state.tickets_df
            )
            
            st.session_state.top_industries = analyze_top_industries(
                st.session_state.companies_df,
                top_n=3
            )

            progress_bar.progress(100)

            st.success("✅ Audit completed successfully!")

    # STEP 6: Display Results
    if st.session_state.audit_results:
        st.markdown("---")
        st.header("📊 Step 6: Audit Results")

        results = st.session_state.audit_results

        # Key Metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Contacts", f"{results['total_contacts']:,}")
        with col2:
            st.metric("Total Companies", f"{results['total_companies']:,}")
        with col3:
            st.metric("Total Tickets", f"{results['total_tickets']:,}")
        with col4:
            total_dups = sum(results['duplicates'].values())
            st.metric("Total Duplicates", f"{total_dups:,}")
        
        # V6 ADVANCED METRICS
        if any([st.session_state.cold_analysis, st.session_state.email_analysis, 
                st.session_state.churn_analysis, st.session_state.critical_tickets]):
            st.markdown("---")
            st.subheader("📊 Advanced Business Metrics")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.session_state.cold_analysis:
                    cold = st.session_state.cold_analysis
                    st.metric(
                        "Cold Contacts (>90d)",
                        f"{cold.get('cold_pct', 0):.1f}%",
                        delta=f"{cold.get('cold_count', 0)} contacts",
                        delta_color="inverse"
                    )
            
            with col2:
                if st.session_state.email_analysis:
                    email = st.session_state.email_analysis
                    st.metric(
                        "Email Validity",
                        f"{email.get('valid_pct', 0):.1f}%",
                        delta=f"{email.get('b2c_pct', 0):.1f}% B2C",
                        delta_color="off"
                    )
            
            with col3:
                if st.session_state.churn_analysis:
                    churn = st.session_state.churn_analysis
                    st.metric(
                        "Churn Risk Contacts",
                        churn.get('at_risk_count', 0),
                        delta=f"${churn.get('arr_at_risk', 0):,.0f} ARR" if churn.get('arr_at_risk') else None,
                        delta_color="inverse"
                    )
            
            with col4:
                if st.session_state.critical_tickets:
                    crit = st.session_state.critical_tickets
                    st.metric(
                        "Critical Tickets (>48h)",
                        crit.get('critical_count', 0),
                        delta=f"{crit.get('avg_resolution', 0):.1f}h avg" if crit.get('avg_resolution') else None,
                        delta_color="inverse"
                    )

        # Visualizations
        st.markdown("---")
        tab1, tab2, tab3, tab4 = st.tabs(["📊 Overview", "🔍 Duplicates Analysis", "⚡ Performance Metrics", "📋 Recommendations"])

        with tab1:
            st.subheader("Data Distribution")

            # Records by type
            records_data = pd.DataFrame({
                'Type': ['Contacts', 'Companies', 'Tickets'],
                'Count': [results['total_contacts'], results['total_companies'], results['total_tickets']]
            })

            fig = px.bar(
                records_data,
                x='Type',
                y='Count',
                title='Records by Type',
                color='Count',
                color_continuous_scale=['#DAA520', '#CD7F32', '#8B4513'],
                text='Count'
            )
            fig.update_traces(texttemplate='%{text:,}', textposition='outside')
            fig = create_powerbi_chart(fig, 'Records by Type')
            st.plotly_chart(fig, use_container_width=True)
            
            add_chart_legend("""
            Shows how many records you have per category.
            <br><br>
            <span style="color: white;">👤 CONTACTS:</span> People in your database
            <br><br>
            <span style="color: white;">🏢 COMPANIES:</span> Organizations or businesses
            <br><br>
            <span style="color: white;">🎫 TICKETS:</span> Support requests
            <br><br>
            <span style="color: white;">📊 IDEAL RATIO:</span><br>
            2-5 contacts per company<br>
            1-3 tickets per contact
            """)

            # Missing data distribution
            if results['missing_data']:
                missing_data = pd.DataFrame({
                    'Type': list(results['missing_data'].keys()),
                    'Missing Values': list(results['missing_data'].values())
                })

                fig = px.pie(
                    missing_data,
                    values='Missing Values',
                    names='Type',
                    title='Missing Data Distribution',
                    color_discrete_sequence=['#CD7F32', '#B8860B', '#DAA520']
                )
                fig = create_powerbi_chart(fig, 'Missing Data Distribution')
                st.plotly_chart(fig, use_container_width=True)
                
                add_chart_legend("""
                Shows where data is missing across your CRM.

🔴 LARGER SLICE:
More missing data in that object type

⚠️ IMPACT:
Missing data blocks automation and reduces insights

📌 PRIORITY:
Focus on filling gaps in your most-used objects first

🎯 TARGET:
Aim for less than 5% missing data per object
                """)

        with tab2:
            st.subheader("Duplicate Records Analysis")

            if results['duplicates']:
                dup_data = pd.DataFrame({
                    'Type': list(results['duplicates'].keys()),
                    'Duplicates': list(results['duplicates'].values())
                })

                fig = px.bar(
                    dup_data,
                    x='Type',
                    y='Duplicates',
                    title='Duplicates by Type',
                    color='Duplicates',
                    color_continuous_scale=['#DAA520', '#CD7F32', '#8B0000'],
                    text='Duplicates'
                )
                fig.update_traces(texttemplate='%{text:,}', textposition='outside')
                fig = create_powerbi_chart(fig, 'Duplicates by Type')
                st.plotly_chart(fig, use_container_width=True)
                
                add_chart_legend("""
                Duplicate records found in your CRM:
                Color intensity: Darker red = more duplicates (higher severity)
                Business impact: Each duplicate costs ~$500/year in lost productivity
                Common causes: Multiple imports, manual entry, lack of validation
                Action: Prioritize merging duplicates in objects with highest count
                """)
            else:
                st.success("✅ No duplicates detected!")


        with tab3:
            st.subheader("⚡ Performance Metrics")
            
            # FREE VERSION DISCLAIMER
            if DEMO_MODE:
                st.warning(f"""
⚠️ **FREE VERSION ANALYSIS**  
Results based on {MAX_ROWS_DEMO}-row sample per file.  
Upgrade to PRO for statistically significant analysis on unlimited data.
""")
            
            # DATA COMPLETENESS
            st.markdown("### 📊 Data Completeness")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.session_state.tickets_completeness:
                    comp = st.session_state.tickets_completeness
                    st.metric(
                        "Tickets Completeness",
                        f"{comp.get('completeness_pct', 0):.1f}%",
                        delta=f"{comp.get('total_fields', 0)} fields"
                    )
            
            with col2:
                if st.session_state.companies_completeness:
                    comp = st.session_state.companies_completeness
                    st.metric(
                        "Companies Completeness",
                        f"{comp.get('completeness_pct', 0):.1f}%",
                        delta=f"{comp.get('total_fields', 0)} fields"
                    )
            
            with col3:
                if st.session_state.overall_quality:
                    qual = st.session_state.overall_quality
                    st.metric(
                        "Overall Quality",
                        f"{qual.get('overall_score', 0):.1f}/100",
                        delta="Average across all objects"
                    )
            
            with col4:
                if st.session_state.quality_improvement:
                    imp = st.session_state.quality_improvement
                    delta_value = imp.get('improvement', 0)
                    st.metric(
                        "Quality Improvement",
                        f"+{delta_value:.1f} pts" if delta_value > 0 else f"{delta_value:.1f} pts",
                        delta="Post-aggregation gain",
                        delta_color="normal" if delta_value > 0 else "inverse"
                    )
            
            st.markdown("---")
            
            # TICKETS PERFORMANCE
            st.markdown("### 🎫 Tickets Performance Metrics")
            
            if st.session_state.tickets_performance:
                perf = st.session_state.tickets_performance
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("Open Tickets", f"{perf.get('open_count', 0):,}")
                    st.metric("Closed Tickets", f"{perf.get('closed_count', 0):,}")
                
                with col2:
                    avg_res = perf.get('avg_resolution_hours', 0)
                    if avg_res > 0:
                        st.metric("Avg Resolution Time", f"{avg_res:.1f}h")
                    else:
                        st.metric("Avg Resolution Time", "N/A")
                
                with col3:
                    sla = perf.get('sla_compliance')
                    if sla is not None:
                        st.metric("SLA Compliance", f"{sla:.1f}%")
                    else:
                        st.markdown("""
<div style='background: #F5F5F7; padding: 15px; border-radius: 8px; 
            border-left: 3px solid #CD7F32; text-align: center;'>
    <strong style='color: #1D1D1F;'>SLA Compliance</strong><br>
    <span style='color: #1D1D1F; font-size: 0.9em;'>
    Upgrade to PRO for advanced metrics
    </span>
</div>
""", unsafe_allow_html=True)
                
                st.markdown("---")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    csat = perf.get('csat_score')
                    if csat is not None:
                        st.metric("CSAT Score", f"{csat:.1f}/5")
                    else:
                        st.markdown("""
<div style='background: #F5F5F7; padding: 15px; border-radius: 8px; 
            border-left: 3px solid #CD7F32; text-align: center;'>
    <strong style='color: #1D1D1F;'>CSAT Score</strong><br>
    <span style='color: #1D1D1F; font-size: 0.9em;'>
    Upgrade to PRO for advanced metrics
    </span>
</div>
""", unsafe_allow_html=True)
                
                with col2:
                    nps = perf.get('nps_score')
                    if nps is not None:
                        st.metric("NPS Score", f"{nps:.1f}")
                    else:
                        st.markdown("""
<div style='background: #F5F5F7; padding: 15px; border-radius: 8px; 
            border-left: 3px solid #CD7F32; text-align: center;'>
    <strong style='color: #1D1D1F;'>NPS</strong><br>
    <span style='color: #1D1D1F; font-size: 0.9em;'>
    Upgrade to PRO for advanced metrics
    </span>
</div>
""", unsafe_allow_html=True)
            
            st.markdown("---")
            
            # TOP INDUSTRIES
            st.markdown("### 🏢 Top Industries")
            
            if st.session_state.top_industries:
                industries = st.session_state.top_industries
                
                if industries.get('no_industry_column'):
                    st.info("ℹ️ No industry column found in companies data")
                elif industries.get('top_industries'):
                    industries_data = pd.DataFrame(industries['top_industries'])
                    
                    fig = px.bar(
                        industries_data,
                        x='name',
                        y='percentage',
                        title='Top 3 Industries',
                        text='percentage',
                        color='percentage',
                        color_continuous_scale=['#DAA520', '#CD7F32', '#8B4513']
                    )
                    fig.update_traces(texttemplate='%{text:.1f}%', textposition='outside')
                    fig.update_xaxes(title='Industry')
                    fig.update_yaxes(title='Percentage (%)')
                    fig = create_powerbi_chart(fig, 'Top 3 Industries Distribution')
                    st.plotly_chart(fig, use_container_width=True)
                    
                    st.markdown("**Industry Breakdown:**")
                    for ind in industries['top_industries']:
                        st.write(f"**{ind['name']}**: {ind['count']:,} companies ({ind['percentage']}%)")
                    
                    total = industries.get('total_companies', 0)
                    if DEMO_MODE and total >= MAX_ROWS_DEMO:
                        st.caption(f"⚠️ Analysis limited to {MAX_ROWS_DEMO} companies. Upgrade to PRO for complete industry analysis.")
                else:
                    st.info("ℹ️ No industry data available")


        with tab4:
            st.subheader("Strategic Recommendations")

            for rec in results['recommendations']:
                with st.expander(f"[{rec['priority']}] {rec['category']}"):
                    st.write(f"**Issue:** {rec['issue']}")
                    st.write(f"**Recommended Action:** {rec['action']}")
                    st.write(f"**Expected Impact:** {rec['impact']}")

        # PRO Access Button (centered)
        st.markdown("---")
        st.markdown("""
        <div style='display: flex; justify-content: center; align-items: center; margin: 2rem 0;'>
            <a href="mailto:wbse.consult@gmail.com" 
               class="interactive-btn" 
               style="text-decoration: none; padding: 20px 60px; font-size: 1.2rem;">
               Contact us for PRO access
            </a>
        </div>
        """, unsafe_allow_html=True)



# ==================== FOOTER ====================
st.markdown("---")
st.markdown("""
<div class="footer">
    CRM Data Quality Analysis Report by Stephanie Jupiter Jacca from WBSE
</div>
""", unsafe_allow_html=True)

